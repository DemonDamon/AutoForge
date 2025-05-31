"""
文档格式转换器
将各种格式的文档转换为Markdown格式
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Optional, List, Dict, Any
import logging
from pathlib import Path

# 延迟导入，避免依赖问题
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

logger = logging.getLogger(__name__)


class BaseConverter(ABC):
    """文档转换器基类"""
    
    @abstractmethod
    def convert(self, file_path: str) -> str:
        """将文件转换为Markdown格式"""
        pass
    
    @abstractmethod
    def can_convert(self, file_path: str) -> bool:
        """检查是否能处理该文件"""
        pass


class PDFConverter(BaseConverter):
    """PDF文档转换器"""
    
    def __init__(self):
        if pdfplumber is None:
            raise ImportError("请安装pdfplumber: pip install pdfplumber")
    
    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def convert(self, file_path: str) -> str:
        """将PDF转换为Markdown"""
        markdown_content = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                markdown_content.append(f"# 文档: {os.path.basename(file_path)}\n")
                
                for i, page in enumerate(pdf.pages):
                    markdown_content.append(f"\n## 第 {i+1} 页\n")
                    
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        # 清理文本
                        text = self._clean_text(text)
                        markdown_content.append(text)
                    
                    # 提取表格
                    tables = page.extract_tables()
                    for j, table in enumerate(tables):
                        if table:
                            markdown_content.append(f"\n### 表格 {j+1}\n")
                            markdown_content.append(self._table_to_markdown(table))
        
        except Exception as e:
            logger.error(f"PDF解析错误: {e}")
            raise
        
        return "\n".join(markdown_content)
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余的空白
        text = re.sub(r'\s+', ' ', text)
        # 保留段落结构
        text = re.sub(r'(?<=[.。!！?？])\s+', '\n\n', text)
        return text.strip()
    
    def _table_to_markdown(self, table: List[List[str]]) -> str:
        """将表格转换为Markdown格式"""
        if not table:
            return ""
        
        markdown_lines = []
        
        # 表头
        header = table[0]
        markdown_lines.append("| " + " | ".join(str(cell or "") for cell in header) + " |")
        markdown_lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # 表体
        for row in table[1:]:
            markdown_lines.append("| " + " | ".join(str(cell or "") for cell in row) + " |")
        
        return "\n".join(markdown_lines)


class WordConverter(BaseConverter):
    """Word文档转换器"""
    
    def __init__(self):
        if Document is None:
            raise ImportError("请安装python-docx: pip install python-docx")
    
    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def convert(self, file_path: str) -> str:
        """将Word文档转换为Markdown"""
        markdown_content = []
        
        try:
            doc = Document(file_path)
            markdown_content.append(f"# 文档: {os.path.basename(file_path)}\n")
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 处理标题
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                        markdown_content.append(f"\n{'#' * (level + 1)} {paragraph.text}\n")
                    else:
                        markdown_content.append(f"{paragraph.text}\n")
            
            # 处理表格
            for i, table in enumerate(doc.tables):
                markdown_content.append(f"\n### 表格 {i+1}\n")
                markdown_content.append(self._table_to_markdown(table))
        
        except Exception as e:
            logger.error(f"Word文档解析错误: {e}")
            raise
        
        return "\n".join(markdown_content)
    
    def _table_to_markdown(self, table) -> str:
        """将Word表格转换为Markdown格式"""
        markdown_lines = []
        
        # 获取表格数据
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        
        if not data:
            return ""
        
        # 表头
        header = data[0]
        markdown_lines.append("| " + " | ".join(header) + " |")
        markdown_lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # 表体
        for row in data[1:]:
            markdown_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(markdown_lines)


class ExcelConverter(BaseConverter):
    """Excel文档转换器"""
    
    def __init__(self):
        if pd is None:
            raise ImportError("请安装pandas和openpyxl: pip install pandas openpyxl")
    
    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.xlsx', '.xls', '.csv'))
    
    def convert(self, file_path: str) -> str:
        """将Excel转换为Markdown"""
        markdown_content = []
        
        try:
            markdown_content.append(f"# 文档: {os.path.basename(file_path)}\n")
            
            if file_path.lower().endswith('.csv'):
                df_dict = {'Sheet1': pd.read_csv(file_path)}
            else:
                # 读取所有工作表
                df_dict = pd.read_excel(file_path, sheet_name=None)
            
            for sheet_name, df in df_dict.items():
                markdown_content.append(f"\n## 工作表: {sheet_name}\n")
                
                # 基本信息
                markdown_content.append(f"- 行数: {len(df)}")
                markdown_content.append(f"- 列数: {len(df.columns)}")
                markdown_content.append(f"- 列名: {', '.join(df.columns)}\n")
                
                # 数据预览
                markdown_content.append("### 数据预览\n")
                markdown_content.append(df.head(10).to_markdown(index=False))
                
                # 数据统计
                if len(df) > 10:
                    markdown_content.append("\n### 数据统计\n")
                    markdown_content.append(df.describe().to_markdown())
        
        except Exception as e:
            logger.error(f"Excel解析错误: {e}")
            raise
        
        return "\n".join(markdown_content)


class ImageConverter(BaseConverter):
    """图片转换器（使用OCR或多模态模型）"""
    
    def __init__(self, use_multimodal: bool = True, multimodal_api_key: Optional[str] = None):
        self.use_multimodal = use_multimodal
        self.multimodal_api_key = multimodal_api_key
        
        if not use_multimodal and (Image is None or pytesseract is None):
            raise ImportError("请安装PIL和pytesseract: pip install pillow pytesseract")
    
    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'))
    
    def convert(self, file_path: str) -> str:
        """将图片转换为Markdown描述"""
        markdown_content = [f"# 图片: {os.path.basename(file_path)}\n"]
        
        try:
            if self.use_multimodal:
                # 使用多模态大模型
                description = self._analyze_with_multimodal(file_path)
                markdown_content.append(description)
            else:
                # 使用OCR
                text = self._extract_text_with_ocr(file_path)
                if text:
                    markdown_content.append("## OCR提取的文本\n")
                    markdown_content.append(text)
                else:
                    markdown_content.append("*未能从图片中提取到文本*")
        
        except Exception as e:
            logger.error(f"图片解析错误: {e}")
            raise
        
        return "\n".join(markdown_content)
    
    def _extract_text_with_ocr(self, file_path: str) -> str:
        """使用OCR提取文本"""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return text.strip()
    
    def _analyze_with_multimodal(self, file_path: str) -> str:
        """使用多模态模型分析图片"""
        # 这里需要集成具体的多模态模型API
        # 暂时返回占位符
        return f"""## 图片分析

![{os.path.basename(file_path)}]({file_path})

### 图片描述
*需要配置多模态模型API来自动生成图片描述*

### 关键信息提取
*需要配置多模态模型API来提取关键信息*
"""


class MarkdownConverter(BaseConverter):
    """Markdown文档转换器（直接返回）"""
    
    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.md', '.markdown'))
    
    def convert(self, file_path: str) -> str:
        """读取Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return f"# 文档: {os.path.basename(file_path)}\n\n{content}" 